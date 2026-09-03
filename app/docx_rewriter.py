# This file rewrites paragraphs after humanization.
# In simple terms: it replaces the old sentence text with the new humanized text,
# keeping the same font (name, size, bold etc.) and removing * markers only
# from sentences that were successfully humanized.

from dataclasses import dataclass

from docx.oxml.ns import qn
from docx.shared import Length, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.docx_engine import IGNORED_SHADING_FILLS


@dataclass(frozen=True)
class FontStyleSnapshot:
    name: str | None = None
    size: Length | None = None
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    color_rgb: RGBColor | None = None


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    # Snapshot font from first run (asterisk markers are textual, so any run will do)
    snapshot = _snapshot_first_marked_font(paragraph)
    if snapshot is None and paragraph.runs:
        snapshot = _snapshot_font(paragraph.runs[0])

    element = paragraph._p
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)

    if not new_text:
        return
    new_run = paragraph.add_run(new_text)
    if snapshot is not None:
        _apply_font_snapshot(new_run, snapshot)
    _strip_marking(new_run)


def replace_paragraph_sentences(
    paragraph: Paragraph,
    sentence_texts: list[str],
    sentence_has_replacement: list[bool] | None = None,
) -> None:
    """
    Sentence-level rewriter. Replaces paragraph content with provided sentence_texts
    (one per segmented sentence, in document order). For successfully humanized
    sentences the paired markers are already removed by the caller;
    failed/skipped sentences keep their original markers.
    Formatting is preserved per-sentence using the first run overlapping that span.
    """
    if not sentence_texts:
        return

    from app.sentence import segment_sentences
    from app.docx_engine import get_asterisk_ranges  # local import to avoid cycle
    from app.sentence import SentenceSpan

    full_text = paragraph.text
    spans = segment_sentences(full_text)
    # If segmentation produced different count than provided texts, fall back to paragraph replacement
    if len(spans) != len(sentence_texts):
        # Join with one space; the caller already stripped successful markers.
        replace_paragraph_text(paragraph, " ".join(sentence_texts))
        return

    # Snapshot per sentence: find first run overlapping that span
    # Build asterisk ranges for reference (not strictly needed for font, but kept)
    asterisk_ranges = get_asterisk_ranges(paragraph)
    # Build per-sentence snapshot list
    snapshots: list[FontStyleSnapshot | None] = []
    run_infos: list[tuple[int, int, FontStyleSnapshot, bool]] = []  # start, end, snapshot, is_asterisk
    offset = 0
    for run in iter_all_runs(paragraph):
        text = run.text or ""
        if not text:
            continue
        # Determine if this run overlaps an asterisk span (textual check approximated)
        is_asterisk = False
        run_start = offset
        run_end = offset + len(text)
        for hs, he in asterisk_ranges:
            if hs < run_end and he > run_start:
                is_asterisk = True
                break
        snap = _snapshot_font(run)
        run_infos.append((run_start, run_end, snap, is_asterisk))
        offset += len(text)

    def snapshot_for_span(span: SentenceSpan) -> FontStyleSnapshot | None:
        for rs, re, snap, is_asterisk in run_infos:
            if rs < span.end and re > span.start:
                return snap
        if run_infos:
            return run_infos[0][2]
        return None

    for span in spans:
        snapshots.append(snapshot_for_span(span))

    element = paragraph._p
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)

    for idx, text in enumerate(sentence_texts):
        if idx > 0:
            sep_run = paragraph.add_run(" ")
            if snapshots[idx] is not None:
                _apply_font_snapshot(sep_run, snapshots[idx])
            _strip_marking(sep_run)
        run = paragraph.add_run(text)
        snap = snapshots[idx]
        if snap is not None:
            _apply_font_snapshot(run, snap)
        _strip_marking(run)


def iter_all_runs(paragraph: Paragraph):
    for item in paragraph.iter_inner_content():
        if isinstance(item, Run):
            yield item
        else:
            yield from item.runs


def _run_is_marked(run: Run) -> bool:
    # Legacy highlight check — not used for asterisk path but kept for compat
    if run.font.highlight_color is not None:
        return True
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return False
    for node in rpr.findall(qn("w:shd")):
        fill = (node.get(qn("w:fill")) or "").lower()
        if fill not in IGNORED_SHADING_FILLS:
            return True
    return False


def _snapshot_first_marked_font(paragraph: Paragraph) -> FontStyleSnapshot | None:
    for run in iter_all_runs(paragraph):
        if _run_is_marked(run):
            return _snapshot_font(run)
    return None


def _snapshot_font(run: Run) -> FontStyleSnapshot:
    font = run.font
    color_rgb = None
    try:
        if font.color is not None and font.color.type is not None:
            color_rgb = font.color.rgb
    except (AttributeError, TypeError):
        color_rgb = None
    return FontStyleSnapshot(
        name=font.name,
        size=font.size,
        bold=font.bold,
        italic=font.italic,
        underline=font.underline,
        color_rgb=color_rgb,
    )


def _strip_marking(run: Run) -> None:
    """Remove legacy highlight/shading (harmless for asterisk mode)."""
    run.font.highlight_color = None
    rpr = run._element.find(qn("w:rPr"))
    if rpr is not None:
        for node in list(rpr.findall(qn("w:shd"))):
            fill = (node.get(qn("w:fill")) or "").lower()
            if fill not in IGNORED_SHADING_FILLS:
                rpr.remove(node)
            else:
                rpr.remove(node)
        for node in list(rpr.findall(qn("w:highlight"))):
            rpr.remove(node)


def _apply_font_snapshot(run: Run, snapshot: FontStyleSnapshot) -> None:
    font = run.font
    if snapshot.name is not None:
        font.name = snapshot.name
    if snapshot.size is not None:
        font.size = snapshot.size
    if snapshot.bold is not None:
        font.bold = snapshot.bold
    if snapshot.italic is not None:
        font.italic = snapshot.italic
    if snapshot.underline is not None:
        font.underline = snapshot.underline
    if snapshot.color_rgb is not None:
        font.color.rgb = snapshot.color_rgb
