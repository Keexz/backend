# This file finds AI-marked sentences and protected sections.
# In simple terms: it reads each paragraph's plain text, looks for *single-asterisk*
# markers with regex, and maps them to sentences.

import io
import re
from dataclasses import dataclass, field

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph

from app.sentence import SentenceSpan, segment_sentences

# Match *text* only when each delimiter is exactly one asterisk.
# Double and unmatched asterisks are ignored.
ASTERISK_RE = re.compile(r"(?<!\*)\*([^*\r\n]+?)\*(?!\*)")
SINGLE_ASTERISK_RE = re.compile(r"(?<!\*)\*(?!\*)")

EXACT_SECTION_TITLES = {
    "TITLE",
    "TITLE PAGE",
    "DECLARATION",
    "CERTIFICATION",
    "CERTIFICATE",
    "DEDICATION",
    "ACKNOWLEDGEMENT",
    "ACKNOWLEDGEMENTS",
    "ACKNOWLEDGMENT",
    "ACKNOWLEDGMENTS",
    "TABLE OF CONTENTS",
    "CONTENTS",
    "LIST OF TABLES",
    "LIST OF FIGURES",
    "APPENDIX",
    "APPENDICES",
    "BIBLIOGRAPHY",
    "REFERENCE",
    "REFERENCES",
}

PREFIXABLE_SECTION_TITLES = ("APPENDIX", "APPENDICES", "REFERENCES")

HEADING_STYLES = {"heading 1", "heading 2"}

MAX_TITLE_LENGTH = 60

# Kept because the rewriter removes old shading from rewritten runs.
IGNORED_SHADING_FILLS = {"", "auto", "ffffff"}


@dataclass
class ParagraphInfo:
    index: int
    text: str
    style_name: str
    has_highlight: bool
    is_heading: bool
    is_section_title: bool
    is_protected: bool
    protection_reason: str
    classified_by: str = "rules"


@dataclass
class SentenceInfo:
    paragraph_index: int
    sentence_index: int  # index within paragraph
    global_index: int  # index across document
    text: str
    start: int
    end: int
    has_highlight: bool
    is_protected: bool
    protection_reason: str


@dataclass
class DocumentAnalysis:
    paragraphs: list[ParagraphInfo] = field(default_factory=list)
    sentences: list[SentenceInfo] = field(default_factory=list)

    @property
    def humanizable_indices(self) -> list[int]:
        return [
            p.index
            for p in self.paragraphs
            if p.has_highlight and not p.is_protected
        ]

    @property
    def humanizable_sentence_indices(self) -> list[int]:
        return [
            s.global_index
            for s in self.sentences
            if s.has_highlight and not s.is_protected
        ]


def load_document(data: bytes) -> DocumentObject:
    return Document(io.BytesIO(data))


def _normalize(text: str) -> str:
    return " ".join(text.split()).strip(" .:;-").upper()


def _is_heading(paragraph: Paragraph) -> bool:
    style_name = ""
    if paragraph.style is not None:
        style_name = (paragraph.style.name or "").lower()
    return style_name in HEADING_STYLES


def _is_section_title(paragraph: Paragraph) -> bool:
    normalized = _normalize(paragraph.text)
    if not normalized or len(normalized) > MAX_TITLE_LENGTH:
        return False
    if normalized in EXACT_SECTION_TITLES:
        return True
    return any(
        normalized.startswith(prefix + " ") or normalized.startswith(prefix + ":")
        for prefix in PREFIXABLE_SECTION_TITLES
    )


def paragraph_has_highlight(paragraph: Paragraph) -> bool:
    # Keep the old function name so existing callers do not need to change.
    return paragraph_has_asterisk(paragraph)


def paragraph_has_asterisk(paragraph: Paragraph) -> bool:
    # Check the concatenated text so markers split across Word runs still work.
    text = paragraph.text or ""
    return bool(ASTERISK_RE.search(text))


def get_highlight_ranges(paragraph: Paragraph) -> list[tuple[int, int]]:
    """
    Keep the old function name while returning strict single-asterisk ranges.
    """
    return get_asterisk_ranges(paragraph)


def get_asterisk_ranges(paragraph: Paragraph) -> list[tuple[int, int]]:
    """
    In simple terms: read the whole paragraph text, find every *...* pair,
    and return where each pair starts and ends.
    """
    text = paragraph.text or ""
    return [(m.start(), m.end()) for m in ASTERISK_RE.finditer(text)]


def strip_asterisk_markers(text: str) -> str:
    """
    Remove paired single markers but keep inner content. *foo* becomes foo.
    Double and unmatched asterisks stay unchanged.
    """
    return ASTERISK_RE.sub(r"\1", text)


def strip_candidate_marker_characters(text: str) -> str:
    """Remove single marker characters from a sentence already known to be marked."""
    return SINGLE_ASTERISK_RE.sub("", text)


def _sentence_has_highlight(span: SentenceSpan, highlight_ranges: list[tuple[int, int]]) -> bool:
    for hs, he in highlight_ranges:
        if hs < span.end and he > span.start:
            return True
    return False


def _sentence_has_asterisk(span: SentenceSpan, asterisk_ranges: list[tuple[int, int]]) -> bool:
    return _sentence_has_highlight(span, asterisk_ranges)


def build_sentence_infos(analysis: DocumentAnalysis, document: DocumentObject) -> None:
    """Populate analysis.sentences based on current paragraphs and document."""
    analysis.sentences.clear()
    global_idx = 0
    for info in analysis.paragraphs:
        paragraph = document.paragraphs[info.index]
        full_text = paragraph.text
        if not full_text.strip():
            continue
        spans = segment_sentences(full_text)
        if not spans:
            continue
        highlight_ranges = get_highlight_ranges(paragraph)
        for s_idx, span in enumerate(spans):
            has_hl = _sentence_has_asterisk(span, highlight_ranges)
            analysis.sentences.append(
                SentenceInfo(
                    paragraph_index=info.index,
                    sentence_index=s_idx,
                    global_index=global_idx,
                    text=span.text,
                    start=span.start,
                    end=span.end,
                    has_highlight=has_hl,
                    is_protected=info.is_protected,
                    protection_reason=info.protection_reason,
                )
            )
            global_idx += 1


def resolve_protection(paragraphs: list[ParagraphInfo]) -> None:
    seen_boundary = False
    in_protected_scope = False

    for info in paragraphs:
        info.protection_reason = ""
        if info.is_section_title:
            seen_boundary = True
            in_protected_scope = True
            info.protection_reason = "section-title"
        elif info.is_heading:
            seen_boundary = True
            in_protected_scope = False
            info.protection_reason = "heading"
        elif in_protected_scope:
            info.protection_reason = "protected-section"
        elif not seen_boundary:
            info.protection_reason = "front-matter"
        info.is_protected = bool(info.protection_reason)


def analyze_document(document: DocumentObject) -> DocumentAnalysis:
    analysis = DocumentAnalysis()

    for index, paragraph in enumerate(document.paragraphs):
        analysis.paragraphs.append(
            ParagraphInfo(
                index=index,
                text=paragraph.text,
                style_name=paragraph.style.name if paragraph.style else "",
                has_highlight=paragraph_has_highlight(paragraph),
                is_heading=_is_heading(paragraph),
                is_section_title=_is_section_title(paragraph),
                is_protected=False,
                protection_reason="",
            )
        )

    resolve_protection(analysis.paragraphs)
    build_sentence_infos(analysis, document)
    return analysis


def rebuild_sentence_infos(analysis: DocumentAnalysis, document: DocumentObject) -> None:
    """Rebuild sentence infos after re-classifying boundaries."""
    build_sentence_infos(analysis, document)
