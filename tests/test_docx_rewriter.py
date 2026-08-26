from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from app.docx_engine import load_document, paragraph_has_highlight
from app.docx_rewriter import replace_paragraph_text

NEW_TEXT = "This sentence was rewritten by a human editor."


def _load(data):
    return load_document(data)


def test_replaces_text_and_strips_highlight(build_docx):
    data = build_docx([{"text": "Old AI text.", "highlighted": True}])
    document = _load(data)
    paragraph = document.paragraphs[0]

    replace_paragraph_text(paragraph, NEW_TEXT)

    assert paragraph.text == NEW_TEXT
    assert paragraph_has_highlight(paragraph) is False


def test_preserves_font_attributes(build_docx):
    data = build_docx(
        [
            {
                "text": "Styled AI text.",
                "highlighted": True,
                "font_name": "Times New Roman",
                "size_pt": 12,
                "bold": True,
                "italic": True,
            }
        ]
    )
    document = _load(data)
    paragraph = document.paragraphs[0]

    replace_paragraph_text(paragraph, NEW_TEXT)

    run = paragraph.runs[0]
    assert run.font.name == "Times New Roman"
    assert run.font.size.pt == 12
    assert run.font.bold is True
    assert run.font.italic is True
    assert run.font.highlight_color is None


def test_prefers_first_highlighted_run_formatting(build_docx):
    data = build_docx(
        [
            {
                "runs": [
                    {
                        "text": "Plain lead-in. ",
                        "font_name": "Arial",
                        "size_pt": 10,
                    },
                    {
                        "text": "highlighted continuation.",
                        "highlighted": True,
                        "font_name": "Georgia",
                        "size_pt": 14,
                    },
                ]
            }
        ]
    )
    document = _load(data)
    paragraph = document.paragraphs[0]

    replace_paragraph_text(paragraph, NEW_TEXT)

    run = paragraph.runs[0]
    assert run.font.name == "Georgia"
    assert run.font.size.pt == 14


def test_paragraph_properties_survive(build_docx):
    data = build_docx([{"text": "Centered AI text.", "highlighted": True}])
    document = _load(data)
    paragraph = document.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    replace_paragraph_text(paragraph, NEW_TEXT)

    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_sibling_paragraphs_remain_byte_identical(build_docx):
    data = build_docx(
        [
            {"text": "Untouched neighbor.", "highlighted": True},
            {"text": "Target AI paragraph.", "highlighted": True},
        ]
    )
    original_document = _load(data)
    untouched_xml_before = original_document.paragraphs[0]._p.xml

    document = _load(data)
    replace_paragraph_text(document.paragraphs[1], NEW_TEXT)

    assert document.paragraphs[0]._p.xml == untouched_xml_before
    assert document.paragraphs[0].text == "Untouched neighbor."


def test_hyperlink_wrapped_runs_are_replaced():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph._p.append(
        parse_xml(
            f'<w:hyperlink {nsdecls("w")}>'
            "<w:r><w:t>old linked text</w:t></w:r>"
            "</w:hyperlink>"
        )
    )

    replace_paragraph_text(paragraph, NEW_TEXT)

    assert paragraph.text == NEW_TEXT
    assert len(paragraph._p.findall(f".//{qn('w:hyperlink')}")) == 0


def test_empty_new_text_clears_paragraph(build_docx):
    data = build_docx([{"text": "Something.", "highlighted": True}])
    document = _load(data)
    paragraph = document.paragraphs[0]

    replace_paragraph_text(paragraph, "")

    assert paragraph.text == ""
    assert paragraph.runs == []


def test_shaded_run_formatting_preserved_and_marking_cleared(build_docx):
    data = build_docx(
        [
            {
                "text": "Shaded AI sentence.",
                "shade_fill": "B9E8F0",
                "font_name": "Georgia",
                "size_pt": 12,
            }
        ]
    )
    document = _load(data)
    paragraph = document.paragraphs[0]
    assert paragraph_has_highlight(paragraph) is True

    replace_paragraph_text(paragraph, NEW_TEXT)

    run = paragraph.runs[0]
    assert paragraph.text == NEW_TEXT
    assert run.font.name == "Georgia"
    assert run.font.size.pt == 12
    assert len(paragraph._p.findall(f".//{qn('w:shd')}")) == 0
    assert paragraph_has_highlight(paragraph) is False


def test_plain_unstyled_paragraph_gets_clean_default_run(build_docx):
    data = build_docx([{"text": "No highlight here."}])
    document = _load(data)
    paragraph = document.paragraphs[0]

    replace_paragraph_text(paragraph, NEW_TEXT)

    assert paragraph.text == NEW_TEXT
    assert paragraph.runs[0].font.name is None
