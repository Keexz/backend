from io import BytesIO

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt


def _set_raw_highlight(run, val):
    rpr = run._element.get_or_add_rPr()
    rpr.append(parse_xml(f'<w:highlight {nsdecls("w")} w:val="{val}"/>'))


def _set_shading(run, fill):
    rpr = run._element.get_or_add_rPr()
    rpr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill}"/>'))


@pytest.fixture
def build_docx():
    def _build(specs):
        document = Document()
        for spec in specs:
            if "runs" in spec:
                runs = spec["runs"]
                if not any(isinstance(r, dict) for r in runs):
                    runs = [
                        r if isinstance(r, dict) else {"text": r[0], "highlighted": r[1]}
                        for r in runs
                    ]
            else:
                runs = [{k: v for k, v in spec.items() if k != "style"}]
            paragraph = document.add_paragraph(style=spec.get("style"))
            for run_spec in runs:
                run = paragraph.add_run(run_spec.get("text", ""))
                font = run.font
                if run_spec.get("highlighted"):
                    font.highlight_color = WD_COLOR_INDEX.BLUE
                if run_spec.get("highlight_val"):
                    _set_raw_highlight(run, run_spec["highlight_val"])
                if run_spec.get("shade_fill"):
                    _set_shading(run, run_spec["shade_fill"])
                if run_spec.get("font_name"):
                    font.name = run_spec["font_name"]
                if run_spec.get("size_pt"):
                    font.size = Pt(run_spec["size_pt"])
                if "bold" in run_spec:
                    font.bold = run_spec["bold"]
                if "italic" in run_spec:
                    font.italic = run_spec["italic"]
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    return _build


@pytest.fixture
def build_docx_with_table():
    def _build(cell_text, highlighted=True):
        document = Document()
        table = document.add_table(rows=1, cols=1)
        run = table.cell(0, 0).paragraphs[0].add_run(cell_text)
        if highlighted:
            run.font.highlight_color = WD_COLOR_INDEX.BLUE
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    return _build
