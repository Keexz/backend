import json

import httpx
from docx import Document
from fastapi.testclient import TestClient

import app.jobs_service as jobs_service
import app.routers.jobs as jobs_router
from app.docx_engine import load_document, paragraph_has_asterisk
from app.job_store import JobStore
from app.main import app

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

THESIS_SPECS = [
    {"text": "THE TITLE OF THE THESIS"},
    {"text": "DECLARATION"},
    {"text": "*I declare this work is original.*"},
    {"text": "CHAPTER ONE", "style": "Heading 1"},
    {"text": "*First AI paragraph.*"},
    {"text": "*Second AI paragraph.*"},
]


def _ryne_ok(content="Humanized body text.", ai_score=5):
    def handler(request):
        return httpx.Response(200, json={"content": content, "aiScore": ai_score})

    return httpx.MockTransport(handler)


def _ryne_always_fail():
    def handler(request):
        return httpx.Response(500, json={"error": "down"})

    return httpx.MockTransport(handler)


def _groq_boundaries(boundaries):
    requests = []

    def handler(request):
        requests.append(httpx.Request)
        payload = request.read()
        import json

        requests[-1] = json.loads(payload)
        content = '{"boundaries": [%s]}' % ", ".join(map(str, boundaries))
        return httpx.Response(
            200,
            json={"choices": [{"message": {"role": "assistant", "content": content}}]},
        )

    return httpx.MockTransport(handler), requests


def _install(monkeypatch, store, ryne=None, groq=None):
    monkeypatch.setattr(jobs_router, "job_store", store)
    if ryne is not None:
        monkeypatch.setattr(jobs_service, "_ryne_transport", lambda: ryne)
    if groq is not None:
        monkeypatch.setattr(jobs_service, "_groq_transport", lambda: groq)
    monkeypatch.setattr(
        jobs_service, "_ryne_backoff", lambda: (0.0, 0.0, 0.0)
    )
    monkeypatch.setattr(jobs_service, "_check_cpu_and_maybe_exit", lambda: None)


def _upload(build_docx, specs=THESIS_SPECS, filename="thesis.docx"):
    return TestClient(app).post(
        "/api/jobs",
        files={"file": (filename, build_docx(specs), DOCX_MIME)},
    )


def test_full_job_lifecycle(monkeypatch, build_docx):
    _install(monkeypatch, JobStore(), ryne=_ryne_ok())
    client = TestClient(app)

    response = _upload(build_docx)
    assert response.status_code == 202
    job_id = response.json()["job_id"]

    status = client.get(f"/api/jobs/{job_id}")
    assert status.status_code == 200
    body = status.json()
    assert body["status"] == "completed"
    assert body["total_paragraphs"] == 2
    assert body["processed_paragraphs"] == 2
    assert {entry["index"] for entry in body["ai_scores"]} == {4, 5}
    assert all(entry["ai_score"] == 5 for entry in body["ai_scores"])
    assert body["failures"] == []

    download = client.get(f"/api/jobs/{job_id}/download")
    assert download.status_code == 200
    assert download.headers["content-type"] == DOCX_MIME

    document = load_document(download.content)
    texts = [p.text for p in document.paragraphs]
    assert texts[0] == "THE TITLE OF THE THESIS"
    assert texts[2] == "*I declare this work is original.*"
    assert texts[4] == "Humanized body text."
    assert texts[5] == "Humanized body text."
    assert not paragraph_has_asterisk(document.paragraphs[4])
    assert not paragraph_has_asterisk(document.paragraphs[5])
    assert paragraph_has_asterisk(document.paragraphs[2]) is True


def test_marker_pair_spanning_sentences_is_sent_as_clean_sentences(monkeypatch, build_docx):
    received_texts = []

    def handler(request):
        received_texts.append(json.loads(request.read())["text"])
        return httpx.Response(200, json={"content": "Humanized.", "aiScore": 0})

    _install(
        monkeypatch,
        JobStore(),
        ryne=httpx.MockTransport(handler),
    )
    specs = [
        {"text": "CHAPTER ONE", "style": "Heading 1"},
        {"text": "*First marked sentence. Second marked sentence.*"},
    ]

    response = _upload(build_docx, specs=specs)
    job_id = response.json()["job_id"]
    body = TestClient(app).get(f"/api/jobs/{job_id}").json()

    assert body["total_sentences"] == 2
    assert received_texts == [
        "First marked sentence.",
        "Second marked sentence.",
    ]


def test_upload_rejects_non_docx_files(monkeypatch, build_docx):
    _install(monkeypatch, JobStore())
    client = TestClient(app)

    response = client.post(
        "/api/jobs",
        files={"file": ("notes.txt", b"plain text", "text/plain")},
    )
    assert response.status_code == 400
    assert response.json()["detail"] == "Only .docx files are supported."


def test_corrupted_docx_marks_job_failed(monkeypatch, build_docx):
    _install(monkeypatch, JobStore())
    client = TestClient(app)

    response = client.post(
        "/api/jobs",
        files={"file": ("broken.docx", b"definitely not a zip", DOCX_MIME)},
    )
    assert response.status_code == 202
    job_id = response.json()["job_id"]

    body = client.get(f"/api/jobs/{job_id}").json()
    assert body["status"] == "failed"
    assert body["error"] == "Uploaded file is not a valid .docx document."

    assert client.get(f"/api/jobs/{job_id}/download").status_code == 409


def test_ryne_failures_reported_and_text_preserved(monkeypatch, build_docx):
    _install(monkeypatch, JobStore(), ryne=_ryne_always_fail())
    client = TestClient(app)

    response = _upload(build_docx)
    job_id = response.json()["job_id"]

    body = client.get(f"/api/jobs/{job_id}").json()
    assert body["status"] == "completed_with_failures"
    assert len(body["failures"]) == 2
    assert all("failed after 4 attempts" in f["error"] for f in body["failures"])

    download = client.get(f"/api/jobs/{job_id}/download")
    document = load_document(download.content)
    assert document.paragraphs[4].text == "*First AI paragraph.*"
    assert paragraph_has_asterisk(document.paragraphs[4]) is True


def test_unknown_job_returns_404(monkeypatch):
    _install(monkeypatch, JobStore())
    client = TestClient(app)

    assert client.get("/api/jobs/does-not-exist").status_code == 404
    assert client.get("/api/jobs/does-not-exist/download").status_code == 404


def test_download_before_ready_conflicts(monkeypatch):
    store = JobStore()
    _install(monkeypatch, store)
    job = store.create("pending.docx")
    client = TestClient(app)

    response = client.get(f"/api/jobs/{job.id}/download")
    assert response.status_code == 409


def test_groq_transport_is_not_used_for_candidates(monkeypatch, build_docx):
    # Verify marked text is humanized while Groq remains unused.
    groq_transport, groq_requests = _groq_boundaries([])
    _install(monkeypatch, JobStore(), ryne=_ryne_ok(), groq=groq_transport)
    client = TestClient(app)

    specs = [
        {"text": "CHAPTER ONE", "style": "Heading 1"},
        {"text": "Intro."},
        {"text": "PREFACE"},
        {"text": "*Marked narrative.*"},
    ]
    response = _upload(build_docx, specs=specs)
    job_id = response.json()["job_id"]

    body = client.get(f"/api/jobs/{job_id}").json()
    assert body["status"] == "completed"
    # Without Groq, PREFACE is not a protected section-title, so the marker is eligible.
    assert body["total_paragraphs"] == 1

    assert len(groq_requests) == 0


def test_document_without_markers_completes_with_zero_work(monkeypatch, build_docx):
    _install(monkeypatch, JobStore(), ryne=_ryne_ok())
    client = TestClient(app)

    specs = [
        {"text": "CHAPTER ONE", "style": "Heading 1"},
        {"text": "Nothing marked here."},
    ]
    response = _upload(build_docx, specs=specs)
    job_id = response.json()["job_id"]

    body = client.get(f"/api/jobs/{job_id}").json()
    assert body["status"] == "completed"
    assert body["total_paragraphs"] == 0
    assert body["processed_paragraphs"] == 0
